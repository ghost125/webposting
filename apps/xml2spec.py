# -*-coding: utf-8-*-
# @Project:posting
# @File   :xml2spec.py
# @Date   :2021-03-18
# @Author  :GDCPL

from lxml import etree
import pandas as pd
from docx import Document
from docx.oxml.shared import OxmlElement,qn
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

from docx.enum.section import WD_ORIENTATION


class xmlparse(object):
    def __init__(self,inxml):
        self.root=etree.parse(inxml).getroot()
        self.s_map={'n':'n','arithmetic':'mean','standardDeviation':'sd','median':'mean','leastSquaresMean':'mean',
                   'geometricMean':'mean','logMean':'mean','number':'mean','leastSquares':'mean',
                    'Ratio of least squares means':'mean','interquartile':'sd/max'}

    def get_study(self):
        return self.root.xpath('trialInformation/sponsorProtocolCode')[0].text


    def get_population_agegroup(self):
        agegroup=self.root.xpath("//trialInformation/populationAgeGroup")[0]
        temp=[]
        for et_i, et in enumerate(agegroup.getchildren()):
            temp.append([et_i+1,et.tag])
        df=pd.DataFrame(temp,columns=['agegroup','category'])
        df['total']='xx'
        return df

    def get_discon_reason(self):
        reason_ids=self.root.xpath('subjectDisposition/reasonsNotCompleted/reasonNotCompleted')
        all_reasons=[]
        for reason in reason_ids:
            _reas=[999,'']
            _reas.append(reason.attrib['id'])
            if reason.xpath('type/value'):
                _reas.append(reason.xpath('type/value')[0].text.split('.')[1])
            else:
                _reas.append('')
            if reason.xpath('otherReason'):
                _reas.append(reason.xpath('otherReason')[0].text)
            else:
                _reas.append('')
            all_reasons.append(_reas[:])
        return all_reasons


    def get_disposition(self):
        periods=self.root.xpath("subjectDisposition/postAssignmentPeriods/postAssignmentPeriod")
        all_periods=[]
        all_groups=[]
        arms=0
        for period_i,period in enumerate(periods):
            arms=max(arms,len(period.xpath('arms/arm')))
            period_id=period.attrib['id']
            if period.xpath('completedMilestone') :
                all_periods.append([period_i+1,period_id,period.xpath('completedMilestone')[0].attrib['id'],
                                    'complete treatment'])
            if period.xpath('startedMilestone') :
                all_periods.append([period_i+1,period_id,period.xpath('startedMilestone')[0].attrib['id'],
                                    'start treatment'])
            if period.xpath('otherMilestones') :
                for o_period in period.xpath('otherMilestones/otherMilestone'):
                    all_periods.append([period_i+1,period_id,o_period.attrib['id'],o_period.xpath('title')[0].text])
            for grp in period.xpath('arms/arm'):
                _temp = [period_i+1]
                for i in grp.attrib:
                    _temp.append(grp.attrib[i])
                if grp.xpath('description'):
                    _temp.append(grp.xpath('description')[0].text)
                else:
                    _temp.append('')
                if grp.xpath('title'):
                    _temp.append(grp.xpath('title')[0].text)
                else:
                    _temp.append('')

                all_groups.append(_temp[:])
        dis_reasons=self.get_discon_reason()
        for reason in dis_reasons:
            all_periods.append(reason)
        max_columns=max([len(i) for i in all_periods])
        collist=['period','periodid','categoryid','details','other']
        if max_columns<=len(collist):
            cols=collist[:(max_columns)]
        df=pd.DataFrame(all_periods,columns=cols)
        for i in range(arms):
            df['group%s' %i]='xx'
        group_df=pd.DataFrame(all_groups,columns=['period','id','description','title'])
        return df,group_df

    def get_category(self,element):
        all_cats=[]
        for _cat in element.xpath('categories/category'):
            _all_cat = []
            _all_cat.append(_cat.attrib['id'])
            _all_cat.append(_cat.xpath('name')[0].text)
            all_cats.append(_all_cat[:])
        return all_cats

    def get_baseline_groups(self):
        all_grps = []
        cols = list(self.root.xpath('//baselineReportingGroups/baselineReportingGroup')[0].attrib.keys())
        cols.insert(0, 'arm order')
        for grpi,grp in enumerate(self.root.xpath('//baselineReportingGroups/baselineReportingGroup')):
            _temp = [grpi]
            for i in grp.attrib:
                _temp.append(grp.attrib[i])
            if grp.xpath('description'):
                if 'description' not in cols:
                    cols.append('description')
                _temp.append(grp.xpath('description')[0].text)

            all_grps.append(_temp[:])
        return pd.DataFrame(all_grps, columns=cols)

    def _get_cat_chars(self,elementlist,ord=0):
        outlist=[]
        total_ny=0
        for cat_char in elementlist:
            if self.get_category(cat_char):
                ord+=1
                _cat=cat_char.tag
                _title=cat_char.xpath("title")[0].text
                _unit=cat_char.xpath('unit')[0].text
                _temp=[]
                for cu_cat in self.get_category(cat_char):
                    _temp=[_cat,ord,_title,_unit]
                    _temp.extend(cu_cat)
                    _temp.append('n')
                    outlist.append(_temp[:])
                total_grp = cat_char.xpath('totalBaselineGroup')
                if total_grp is not None:
                    total_ny=1
        return outlist,total_ny,ord

    def _get_cont_chars(self,elementlist,ord=0):
        outlist=[]
        for con_char in elementlist:
            if con_char.xpath('title'):
                ord += 1
                _cat = con_char.tag
                _title = con_char.xpath("title")[0].text
                _unit = con_char.xpath('unit')[0].text
                if con_char.xpath('centralTendencyType/value'):
                    _temp = [_cat, ord, _title, _unit, ' ',
                             con_char.xpath('centralTendencyType/value')[0].text.split('.')[1]]
                    _temp.append('mean')
                    outlist.append(_temp)
                if con_char.xpath('dispersionType/value'):
                    _temp = [_cat, ord, _title, _unit, ' ',
                             con_char.xpath('dispersionType/value')[0].text.split('.')[1]]
                    _temp.append('sd/max')
                    outlist.append(_temp)
        return outlist,ord

    def get_baseline(self):
        baseline_chars=self.root.xpath("//baselineCharacteristics")[0]
        baselinegroups=len(self.root.xpath('baselineCharacteristics/baselineReportingGroups/baselineReportingGroup'))
        total_ny=0
        cat_elements=baseline_chars.xpath("studyCategoricalCharacteristics/studyCategoricalCharacteristic")
        outlist0,total_ny0,ord=self._get_cat_chars(cat_elements,0)
        cont_elements=baseline_chars.xpath("studyContinuousCharacteristics/studyContinuousCharacteristic")
        outlist1,ord=self._get_cont_chars(cont_elements,0)
        gender_elements=baseline_chars.xpath("genderCategoricalCharacteristic")
        outlist2,total_ny2,ord=self._get_cat_chars(gender_elements,0)
        age_cont_elements=baseline_chars.xpath("ageContinuousCharacteristic")
        outlist3,ord=self._get_cont_chars(age_cont_elements,0)
        age_cat_elements=baseline_chars.xpath("ageCategoricalCharacteristic")
        outlist4,total_ny4,ord=self._get_cat_chars(age_cat_elements,0)
        baselines=outlist0+outlist1+outlist2+outlist3+outlist4
        df=pd.DataFrame(baselines,columns=['category','_nr_',"title",'measureunit','categoryid','detail','stats'])
        for i in range(baselinegroups):
            df['group%s' %i]='xx N=xx'
        if max(total_ny,total_ny0,total_ny2,total_ny4)>0:
            df['total']='xx N=xx'
        return df

    def get_baseline2(self):
        baseline_chars=self.root.xpath("//baselineCharacteristics")[0]
        rep_mode=baseline_chars.xpath('reportingModel/value')[0].text.split('.')[1]
        if rep_mode.upper()=='ARMS':
            barms=self.get_baseline_groups()
            bsets=pd.DataFrame()
        else:
            bsets=self.get_outcome_sets()
            barms=pd.DataFrame()

        df=pd.DataFrame()
        total_flag=False
        #studyCategoricalCharacteristics
        bas_characters=['studyCategoricalCharacteristics/studyCategoricalCharacteristic',
                        'studyContinuousCharacteristics/studyContinuousCharacteristic',
                        'ageContinuousCharacteristic','genderCategoricalCharacteristic',
                        'ageCategoricalCharacteristic']
        for basi,bas in enumerate(bas_characters):
            for iscc,scc in enumerate(baseline_chars.xpath(bas)):
                bases = []
                title = scc.xpath('title')[0].text
                ord =basi+1
                descrp = scc.xpath('description')[0].text if scc.xpath('description') else None
                unit = scc.xpath('unit')[0].text if scc.xpath('unit') else None

                ptype = '.'.join(scc.xpath('centralTendencyType/value')[0].text.split('.')[1:]) \
                    if scc.xpath('centralTendencyType/value') else None
                dtype = '.'.join(scc.xpath('dispersionType/value')[0].text.split('.')[1:]) \
                    if scc.xpath('dispersionType/value') else None
                if scc.xpath('categories/category'):
                    all_cats = self.get_category(scc)
                    for cat in all_cats:
                        _temp = [ord,iscc+1, title, descrp, unit, cat[0], cat[1]]
                        if ptype is not None and (dtype is None or dtype == 'na'):
                            _temp.append(ptype)
                            bases.append(_temp[:])
                        elif (ptype is None or ptype == 'na') and dtype is not None:
                            _temp.append(dtype)
                            bases.append(_temp[:])
                        elif (ptype is not None and ptype != 'na') and (dtype is not None and dtype != 'na'):
                            _temp.append(ptype)
                            bases.append(_temp[:])
                            _temp[-1] = dtype
                            bases.append(_temp[:])
                        else:
                            _temp.append('n')
                            bases.append(_temp)

                else:
                    _temp = [ord,iscc+1, title, descrp, unit, ' ', ' ']
                    if ptype is not None and (dtype is None or dtype == 'na'):
                        _temp.append(ptype)
                        bases.append(_temp[:])
                    elif (ptype is None or ptype == 'na') and dtype is not None:
                        _temp.append(dtype)
                        bases.append(_temp[:])
                    elif (ptype is not None and ptype != 'na') and (dtype is not None and dtype != 'na'):
                        _temp.append(ptype)
                        bases.append(_temp[:])
                        _temp[-1] = dtype
                        bases.append(_temp[:])
                    else:
                        _temp.append('n')
                        bases.append(_temp)
                df_temp = pd.DataFrame(bases,
                                       columns=['order', '_nr_', 'title', 'description', 'unit',
                                                'categoryid', 'text', 'stat'])
                df_temp['stats']=df_temp['stat'].map(self.s_map)

                if rep_mode.upper()=='ARMS':
                    c_arms = self.get_baseline_groups()
                    for arm_i, arm in enumerate(scc.xpath('reportingGroups/reportingGroup')):
                        c_arms_dict = c_arms[c_arms['id'] == arm.attrib['baselineReportingGroupId']].to_dict(orient='records')[0]
                        df_temp['group%s' % c_arms_dict['arm order']] = 'xx N=xx'
                    if scc.xpath('totalBaselineGroup'):
                        if scc.xpath('totalBaselineGroup')[0].getchildren():
                            df_temp['total']='xx N=xx'
                            total_flag=True
                if rep_mode.upper()!='ARMS':
                    sets = {i['id']: i['order'] for i in self.get_outcome_sets().to_dict(orient='records')}
                    for set_i, set in enumerate(
                            scc.xpath('subjectAnalysisSets/subjectAnalysisSet')):
                        df_temp['set%s' % sets[set.attrib['subjectAnalysisSetId']]] = 'xx N=xx'
                df = pd.concat([df, df_temp], sort=False)
                df.reset_index(drop=True)
                df.replace('na', ' ', inplace=True)
                df.fillna(' ', inplace=True)
                grps = sorted([i for i in df.columns if i.startswith('group')])
                if total_flag:
                    grps.append('total')
                sets = sorted([i for i in df.columns if i.startswith('set')])
                df = df[['order', '_nr_', 'title', 'description', 'unit',
                        'categoryid', 'text', 'stat', 'stats'] + grps + sets]

        return barms,bsets,df

    def get_outcome_groups(self):
        all_grps = []
        cols = list(self.root.xpath('//armReportingGroups/armReportingGroup')[0].attrib.keys())
        cols.insert(0,'endpoint order')
        cols.insert(1, 'arm order')
        for epi,endpoint in enumerate(self.root.xpath('endPoints/endPoint')):
            for igrp,grp in enumerate(endpoint.xpath('armReportingGroups/armReportingGroup')):
                _temp = [epi+1,igrp]
                for i in grp.attrib:
                    _temp.append(grp.attrib[i])
                if grp.xpath('description'):
                    if 'description' not in cols:
                        cols.append('description')
                    _temp.append(grp.xpath('description')[0].text)
                if grp.xpath('title'):
                    if 'title' not in cols:
                        cols.append('title')
                    _temp.append(grp.xpath('title')[0].text)

                all_grps.append(_temp[:])
        return pd.DataFrame(all_grps, columns=cols)

    def get_outcome_sets(self):
        all_sets = []
        set_cols=['order']
        for iset, set in enumerate(self.root.xpath('subjectAnalysisSets/subjectAnalysisSet')):
            _temp=[iset]
            for key,value in set.attrib.items():
                _temp.append(value)
                if key not in set_cols:
                    set_cols.append(key)
            for et in set.getchildren():
                if et.text:
                    _temp.append(et.text)
                    if et.tag not in set_cols:
                        set_cols.append(et.tag)

            all_sets.append(_temp[:])
        return pd.DataFrame(all_sets, columns=set_cols)


    def get_outcome(self):
        endpoints=self.root.xpath('//endPoints')[0]
        outcomeans=[]
        ord=0
        df=pd.DataFrame()
        for endp in endpoints.xpath('endPoint'):
            outcomes = []
            title=endp.xpath('title')[0].text
            ord+=1
            descrp=endp.xpath('description')[0].text if endp.xpath('description') else None
            countable=endp.xpath('countable')[0].text
            unit=endp.xpath('unit')[0].text if endp.xpath('unit') else None
            timeframe=endp.xpath('timeFrame')[0].text if endp.xpath('description') else None
            type_='.'.join(endp.xpath('type/value')[0].text.split('.')[1:]) if endp.xpath('type/value') else None
            ptype='.'.join(endp.xpath('centralTendencyType/value')[0].text.split('.')[1:]) \
                if endp.xpath('centralTendencyType/value') else None
            dtype='.'.join(endp.xpath('dispersionType/value')[0].text.split('.')[1:]) \
                if endp.xpath('dispersionType/value') else None
            if endp.xpath('categories/category'):
                all_cats=self.get_category(endp)
                for cat in all_cats:
                    _temp=[ord,countable,title,descrp,unit,timeframe,type_,cat[0],cat[1]]
                    if ptype is not None and (dtype is None or dtype=='na'):
                        _temp.append(ptype)
                        outcomes.append(_temp[:])
                    elif (ptype is None or ptype=='na') and dtype is not None:
                        _temp.append(dtype)
                        outcomes.append(_temp[:])
                    elif (ptype is not None and ptype !='na') and (dtype is not None and dtype !='na'):
                        _temp.append(ptype)
                        outcomes.append(_temp[:])
                        _temp[-1]=dtype
                        outcomes.append(_temp[:])
                    else:
                        _temp.append('n')
                        outcomes.append(_temp)

            else:
                _temp=[ord,countable,title,descrp,unit,timeframe,type_,' ',' ']
                if ptype is not None and (dtype is None or dtype == 'na'):
                    _temp.append(ptype)
                    outcomes.append(_temp[:])
                elif (ptype is None or ptype == 'na') and dtype is not None:
                    _temp.append(dtype)
                    outcomes.append(_temp[:])
                elif (ptype is not None and ptype != 'na') and (dtype is not None and dtype != 'na'):
                    _temp.append(ptype)
                    outcomes.append(_temp[:])
                    _temp[-1] = dtype
                    outcomes.append(_temp[:])
                else:
                    _temp.append('n')
                    outcomes.append(_temp)
            df_temp=pd.DataFrame(outcomes,columns=['order','countable','title','description','unit','timeframe','type',
                                               'categoryid','text','stat'])
            df_temp['stats'] = df_temp['stat'].map(self.s_map)
            if endp.xpath('armReportingGroups/armReportingGroup'):
                c_arms=self.get_outcome_groups()
                c_arms=c_arms[c_arms['endpoint order']==ord]
                for arm_i,arm in enumerate(endp.xpath('armReportingGroups/armReportingGroup')):
                    c_arms_dict=c_arms[c_arms['armId']==arm.attrib['armId']].to_dict(orient='records')[0]
                    df_temp['group%s' %c_arms_dict['arm order']]='xx N=xx'
            if endp.xpath('subjectAnalysisSetReportingGroups/subjectAnalysisSetReportingGroup'):
                sets={i['id']:i['order'] for i in self.get_outcome_sets().to_dict(orient='records')}
                for set_i, set in enumerate(endp.xpath('subjectAnalysisSetReportingGroups/subjectAnalysisSetReportingGroup')):
                    df_temp['set%s' % sets[set.attrib['subjectAnalysisSetId']]] = 'xx N=xx'
            df=pd.concat([df,df_temp],sort=False)

            if endp.xpath('statisticalAnalyses/statisticalAnalysis'):
                for endpa in endp.xpath('statisticalAnalyses/statisticalAnalysis'):
                    titlea=endpa.xpath('title')[0].text if endpa.xpath('title') else None
                    descrpa = endpa.xpath('description')[0].text if endpa.xpath('description') else None
                    other_method=endpa.xpath('statisticalHypothesisTest/otherMethod')[0].text if endpa.xpath(
                        'statisticalHypothesisTest/otherMethod') else None
                    othertype=endpa.xpath('parameterEstimate/otherType')[0].text \
                        if endpa.xpath('parameterEstimate/otherType') else None
                    arms=[]
                    for armid in endpa.xpath('armComparisonGroupId'):
                        arms.append(armid.text)
                    sets = []
                    for setid in endpa.xpath('subjectAnalysisSetComparisonGroupId'):
                        sets.append(setid.text)
                    pvalue='xx'
                    if endpa.xpath('parameterEstimate'):
                        for parmE in endpa.xpath('parameterEstimate'):
                            if parmE.xpath('confidenceInterval/lowerLimit'):
                                lowerlimit='xx'
                            else:
                                lowerlimit=''
                            if parmE.xpath('confidenceInterval/percentage'):
                                percentage='xx'
                            else:
                                percentage=''
                            if parmE.xpath('confidenceInterval/upperLimit'):
                                upperlimit='xx'
                            else:
                                upperlimit=''
                            if parmE.xpath('effectEstimate'):
                                effectestimate = 'xx'
                            else:
                                effectestimate = ''
                            if parmE.xpath('variabilityEstimate/dispersionValue'):
                                cv = 'xx'
                            else:
                                cv = ''
                            if parmE.xpath('pointEstimate'):
                                pointestimate='xx'
                            else:
                                pointestimate=''

                    outcomeans.append([ord,titlea,descrpa,other_method,pvalue,percentage,lowerlimit,upperlimit,
                                       effectestimate,othertype,pointestimate,cv,' vs '.join(arms),
                                       ' vs '.join(sets)])
        df.replace('na',' ',inplace=True)
        df.fillna(' ',inplace=True)
        grps = sorted([i for i in df.columns if i.startswith('group')])
        sets = sorted([i for i in df.columns if i.startswith('set')])
        df = df[['order', 'countable', 'title', 'description', 'unit', 'timeframe', 'type',
                 'categoryid', 'text', 'stat','stats'] + grps + sets]

        sdf=pd.DataFrame(outcomeans, columns=['order', 'title', 'description', 'othermethod', 'pvalue', 'percentage',
                                          'lowerlimit', 'upperlimit', 'effectestimate', 'othertype', 'pointestimate',
                                          'variabilityEstimate','armids','setids'])
        sdf.replace('na',' ',inplace=True)

        return df,sdf

    def get_country(self):
        all_countries=self.root.xpath('trialInformation/countrySubjectCounts/countrySubjectCount')
        all_c=[]
        for cntry in all_countries:
            c_cntry=[]
            c_cntry.append(cntry.xpath('country/eutctId')[0].text)
            c_cntry.append(cntry.xpath('country/version')[0].text)
            c_cntry.append('xx')
            all_c.append(c_cntry)
        return pd.DataFrame(all_c,columns=['eutctid','version','count'])



class createSpec(object):
    def __init__(self):
        self.doc=Document()
        self.doc.styles['Normal'].font.name = 'Times New Roman'
        self.doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'),u'宋体')
        current_section=self.doc.sections[-1]
        current_section.orientation=WD_ORIENTATION.LANDSCAPE
        new_width, new_height = current_section.page_height, current_section.page_width
        current_section.page_width = new_width
        current_section.page_height = new_height
        self.doc.add_heading('Posting ERF XML Specification',0)
        self.doc.add_paragraph('''The specification used for prepare the input contents of each section''')
        self.doc.add_paragraph('''NOTE: The maximum count of subjects in this posting XML file are the population used \
for baseline analysis(usually Fas), like by country count,age group, start of treatment and so on.''')
        self.doc.add_paragraph('''NOTE: Except by country count, age group parts, all the other needed information can \
be find in key tables(topline results). ''')
        self.doc.add_paragraph('''NOTE: if more than one group used in corresponding section, the order of treatment group \
variables from the input dataset must same as the arm order of each section.''')
        self.doc.add_paragraph('''NOTE: for continuous characters, only 3 stats is used and their means are:
    mean: arithmetic mean; geometric mean; least squares mean; Log mean; Median; Number
    sd: standard deviation; standrad error; Q1; min; low range of CI or any other pair statistics
    max: Q3; max; high range of CI or any other pair statistics    ''')
        self.doc.add_paragraph('''    for Categorical characteristics, only have one stats: n and means number of subjects ''')

        self.doc.add_paragraph('''NOTE: stats variable must be exist in the input dataset for baseline and endpoints section \
    and its values are:
        for continue variables: mean, sd, max
        for category variables: n''')

    def add_table(self,indf):
        if isinstance(indf,pd.DataFrame):
            tb_header = indf.columns
            table = self.doc.add_table(rows=1, cols=len(tb_header), style="Table Grid")
            table.autofit = True
            header_cells = table.rows[0].cells
            for i in range(len(tb_header)):
                tmp_str = r'<w:shd {} w:fill="' + str.strip('D9D9D9') + r'"/>'
                shading_elm_1 = parse_xml(tmp_str.format(nsdecls('w')))
                # shading_elm_1 = parse_xml(r'<w:shd {} w:fill="6A5ACD"/>'.format(nsdecls('w')))
                header_cells[i]._tc.get_or_add_tcPr().append(shading_elm_1)
                header_cells[i].text = tb_header[i]
                if 0 < i < len(tb_header):
                    tc = header_cells[i]._tc
                    tcPr = tc.get_or_add_tcPr()
                    tcVAlign = OxmlElement('w:vAlign')
                    tcVAlign.set(qn('w:val'), "center")
                    tcPr.append(tcVAlign)
            m, n = indf.shape
            for i in range(m):
                row_cells = table.add_row().cells
                for j in range(n):
                    if isinstance(indf.iloc[i,j],str):
                        row_cells[j].text = indf.iloc[i,j]
                    elif isinstance(indf.iloc[i,j],int):
                        row_cells[j].text = str(indf.iloc[i, j])
                    elif isinstance(indf.iloc[i,j],float):
                        row_cells[j].text = str(indf.iloc[i, j])
                    elif isinstance(indf.iloc[i,j],list):
                        row_cells[j].text = '/'.join(indf.iloc[i, j])
                    elif pd.isna(indf.iloc[i,j]):
                        row_cells[j].text=' '
                    else:
                        row_cells[j].text = str(indf.iloc[i, j])

            self.doc.add_paragraph(' ')

    # trialInformation
    def add_trial_information(self,indf,countrydf=None):
        self.doc.add_paragraph('')
        self.doc.add_heading('Trial Information',level=1)
        if countrydf is not None:
            self.doc.add_heading('Count of patients of each country', level=2)
            self.doc.add_paragraph('The count of patients of each Country are:')
            if countrydf.shape[0]>0:
                self.add_table(countrydf)
            else:
                countrydf=pd.DataFrame([['100000000403',	'12',	'xx'],
['100000000362',	'12',	'xx'],
['100000000557',	'12',	'xx']],columns=['eutctid',	'version',	'count'])
                self.add_table(countrydf)
            self.doc.add_paragraph('In the input dataset, eutcid and version is from the ship excel with this tool, which \
may also be update by the user as needed.')
            self.doc.add_paragraph(
                'for details please visit: https://spor.ema.europa.eu/rmswi/#/lists/100000000002/terms\n')
            self.doc.add_paragraph('count: the count of patients of each country, extract from clinical database')

        self.doc.add_heading('Age category', level=2)
        self.doc.add_paragraph('The category for reporting Age are:')
        self.add_table(indf)
        self.doc.add_paragraph('agegroup: age category, mandatory')
        self.doc.add_paragraph('total: the variable contain the actual result, the variable can be any validate variable name')
        self.doc.add_paragraph('''variable agegroup must exist in the input dataset, and the meaning of its values are:
        1: 'In Utero', 2: 'Pre-term newborn - gestational age < 37 wk', 3: 'Newborns (0-27 days)',  4: 'Infants and toddlers
         (28 days-23 months)', 5: 'Children (2-11 years)',  6: 'Adolescents (12-17 years)', 7: 'Adults (18-64 years)', 
         8: 'Elderly (From 65-84 years)', 9: 'Elderly 85 years and over' ''')

    # subjectDisposition
    def add_subject_disposition(self,arm_indf,indf):
        self.doc.add_paragraph('')
        self.doc.add_heading('Subject Disposition', level=1)
        #self.doc.add_paragraph('Please note the withdraw reason can only be below so remap the withdraw reason in DB is needed:')
        tpdf=pd.DataFrame()
        tpdf['Picklist']=['Adverse event, not fatal','Adverse event, serious fatal','Consent withdrawal by subject',
                         'lack of efficacy', 'Lost to follow up','Physician decision','Pregnancy','Protocol deviation',
                          'Transferred to another arm/group','Other (please specify)']
        #self.add_table(tpdf)
        self.doc.add_paragraph('')

        self.doc.add_heading('The arms of this section',level=2)
        self.add_table(arm_indf)

        self.doc.add_paragraph('')
        self.doc.add_heading('Periods and discontinuation reasons', level=2)
        self.doc.add_paragraph('Shells')
        self.add_table(indf)
        self.doc.add_paragraph('period: flag variable of each period, mandatory')
        self.doc.add_paragraph('periodid: ID of each period, extract from input xml, only for reference')
        self.doc.add_paragraph('categoryid: ID of each category,extract from input xml, mandatory')
        self.doc.add_paragraph('details/other: detail wording of each category, extract from xml and for reference')
        self.doc.add_paragraph('groupxx: Treatment group variables')
        self.doc.add_paragraph('NOTE: Only exist reasons in input dataset will be kept')


    # baseline section
    def add_baseline(self,arm_indf,set_indf,indf):
        self.doc.add_paragraph('')
        self.doc.add_heading('Baseline Characteristics', level=1)
        if not arm_indf.empty:
            self.doc.add_heading('The arms of this section',level=2)
            self.add_table(arm_indf)
        if not set_indf.empty:
            self.doc.add_heading('The analysis sets of this section',level=2)
            self.add_table(arm_indf)

        self.doc.add_paragraph('NOTE: for all shells at below:')
        self.doc.add_paragraph('_nr_: variable output order in each category, mandatory')
        self.doc.add_paragraph('title: title of each variable, extract from input xml, only for reference')
        self.doc.add_paragraph('catetoryid: Extract from xml and mandatory')
        self.doc.add_paragraph('text: Extract from xml and only for reference')
        self.doc.add_paragraph('stat: used to prepare STATS, which may not be completed automatically, mandatory')
        if not arm_indf.empty:
            self.doc.add_paragraph('groupxx/total: Treatment group variables')
        if not set_indf.empty:
            self.doc.add_paragraph('setxx: analysis sets')

        self.doc.add_paragraph('NOTE: For categorical characters, only exist category in input dataset will be kept')

        self.doc.add_heading('Study Categorical characteristics', level=2)
        c_df=indf[indf['order']==1].drop('order',axis=1)
        if len(c_df)>0:
            self.doc.add_paragraph('Shells:')
            self.add_table(c_df)
        else:
            self.doc.add_paragraph("Not Applicable.")

        self.doc.add_heading('Study Continuous characteristics', level=2)
        c_df = indf[indf['order'] == 2].drop('order',axis=1)
        if len(c_df)>0:
            self.doc.add_paragraph('Shells:')
            self.add_table(c_df)
        else:
            self.doc.add_paragraph("Not Applicable.")

        self.doc.add_heading('Age Continuous characteristics', level=2)
        c_df = indf[indf['order'] == 3].drop('order',axis=1)
        if len(c_df)>0:
            self.doc.add_paragraph('Shells:')
            self.add_table(c_df)
        else:
            self.doc.add_paragraph("Not Applicable.")

        self.doc.add_heading('Gender categorical characteristics', level=2)
        c_df = indf[indf['order'] == 4].drop('order',axis=1)
        if len(c_df)>0:
            self.doc.add_paragraph('Shells:')
            self.add_table(c_df)
        else:
            self.doc.add_paragraph("Not Applicable.")

        self.doc.add_heading('Age categorical characteristics', level=2)
        c_df = indf[indf['order'] == 5].drop('order',axis=1)
        if len(c_df)>0:
            self.doc.add_paragraph('Shells:')
            self.add_table(c_df)
        else:
            self.doc.add_paragraph("Not Applicable.")

    # endPoints
    def add_endpoints(self,indf,arm_indf, set_indf, statdf):
        self.doc.add_paragraph('')
        self.doc.add_heading('EndPoints', level=1)
        if not arm_indf.empty:
            self.doc.add_heading('The arms of this section',level=2)
            self.add_table(arm_indf)
        else:
            self.doc.add_paragraph('Not Applicable')
        if not set_indf.empty:
            self.doc.add_heading('The sets of this section',level=2)
            self.add_table(set_indf)
        else:
            self.doc.add_paragraph('Not Applicable')

        self.doc.add_heading('The summary tables are:',level=2)
        self.doc.add_paragraph('Shells:')
        self.add_table(indf)
        self.doc.add_paragraph('order: variable output order in each endpoint, mandatory')
        self.doc.add_paragraph('countable: Bool(true/false) variable to show the variable is continued or characters, mandatory')
        self.doc.add_paragraph(
            'title: title of each variable, extract from input xml, only for reference and add new title as needed')
        self.doc.add_paragraph(
            'description: description of each variable, extract from input xml, only for reference and add new title as needed')
        self.doc.add_paragraph('unit: Extract from xml and only for reference')
        self.doc.add_paragraph('timeframe/type: Extract from xml and only for reference')
        self.doc.add_paragraph('catetoryid: Extract from xml and mandatory')
        self.doc.add_paragraph('text: Extract from xml and only for reference')
        self.doc.add_paragraph('stat: used to prepare STATS, which may not be completed automatically, mandatory')
        if not arm_indf.empty:
            self.doc.add_paragraph('groupxx/total: Treatment group variables')
        if not set_indf.empty:
            self.doc.add_paragraph('setxx: analysis sets')
        self.doc.add_paragraph('NOTE: For categorical characters, only exist category in input dataset will be kept')

        if not statdf.empty:
            self.doc.add_heading('The statistical analysis tables are:', level=2)
            self.doc.add_paragraph('Shells:')
            self.add_table(statdf)
            self.doc.add_paragraph('order: variable used to link with summary part, mandatory')
            self.doc.add_paragraph(
                'title: title of each variable, extract from input xml, only for reference and add new title as needed')
            self.doc.add_paragraph(
                'description: description of each variable, extract from input xml, only for reference and add new title as needed')
            self.doc.add_paragraph('othermethod: Extract from xml and only changed if needed')
            self.doc.add_paragraph('pvalue/percentage/lowerlimit/upperlimit/effectestimate/pointestimate:'
                                   ' input statistical analysis results')
            self.doc.add_paragraph('othertype: Extract from xml and only changed if needed')
            self.doc.add_paragraph('armids: Extract from xml and used to show which groups are used in this statistical analysis')


    def Save(self,filename):
        self.doc.save(filename)




if __name__=='__main__':
    test=xmlparse(r'C:\Studies\posting\pilot\18933\18933_PharmaCM ERF_without value_2021-05-19.xml')

    spec=createSpec()
    overllage=test.get_population_agegroup()
    country=test.get_country()
    spec.add_trial_information(overllage,country)
    disp,group_disp=test.get_disposition()
    spec.add_subject_disposition(group_disp,disp)
    group_baseline,set_baseline,baseline=test.get_baseline2()
    spec.add_baseline(group_baseline,set_baseline,baseline)
    group_endp=test.get_outcome_groups()
    endp,endp_a=test.get_outcome()
    sets = test.get_outcome_sets()
    spec.add_endpoints(endp,group_endp,sets,endp_a)

    spec.Save(r'C:\Studies\posting\pilot\18933\spec.docx')
